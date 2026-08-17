from __future__ import annotations

import hashlib
import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[3]
RULE = ROOT / "deliverables" / "skills" / "daily-work-analysis"
TEMPLATE_SKILL = ROOT / "deliverables" / "skills" / "daily-report-docx-template"
sys.path[:0] = [str(RULE / "scripts"), str(TEMPLATE_SKILL / "scripts")]

from audit_daily_report import audit_daily_report  # noqa: E402
from build_fact_package import build_daily_facts  # noqa: E402
from package_contract import validate_package  # noqa: E402
from render_daily_report import render_daily_report  # noqa: E402
from validate_analysis_result import validate_analysis_result  # noqa: E402
from datetime import date  # noqa: E402


REPORT_DATE = "2026-07-31"
MISSING = ("王达伟", "吴鹏")
LEAVE = ("周菁", "陈苏", "刘茜", "杨瑞", "王海娜")


def canonical(value: object) -> bytes:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")


def leadership_documents() -> dict[str, object]:
    names = [*MISSING, *LEAVE, *(f"测试员工{index:03d}" for index in range(1, 89))]
    roster = [{"person_id": f"p{index:03d}", "name": name, "group": "行一二部",
               "personnel_type": "employee", "effective_start_date": "2026-01-01",
               "effective_end_date": None, "include_in_statistics": True}
              for index, name in enumerate(names, 1)]
    attendance = [{"person_id": f"p{index:03d}", "date": REPORT_DATE,
                   "morning_status": "leave", "afternoon_status": "leave",
                   "leave_type": "annual", "leave_period": "full_day",
                   "source_record_id": f"leave-{index:03d}"} for index in range(3, 8)]
    submitted_ids = [f"p{index:03d}" for index in range(8, 96)]
    reports = [{"report_id": f"report-{person_id}", "person_id": person_id, "date": REPORT_DATE,
                "submission_status": "submitted", "first_submitted_at": f"{REPORT_DATE}T17:00:00+08:00",
                "final_submitted_at": f"{REPORT_DATE}T17:00:00+08:00", "raw_text": "测试日报工作事项",
                "snapshot_id": f"snapshot-report-{person_id}"} for person_id in submitted_ids]
    tasks = []
    relations = []
    for index in range(115):
        person_id = submitted_ids[index % len(submitted_ids)]
        task_id = f"task-{index + 1:03d}"
        status = "completed" if index < 11 else "blocked" if index < 23 else "in_progress"
        tasks.append({"task_id": task_id, "report_id": f"report-{person_id}", "date": REPORT_DATE,
                      "time_period": "morning" if index % 2 == 0 else "afternoon",
                      "raw_text": f"项目推进事项{index + 1}", "task_type": "project",
                      "project_candidate_id": "project-main", "status": status,
                      "output": "形成阶段成果" if status == "completed" else None,
                      "risk": "进度待协调" if status == "blocked" else None,
                      "blocker": "外部反馈待确认" if status == "blocked" else None,
                      "next_step": "继续推进"})
        relations.append({"relation_id": f"relation-{index + 1:03d}", "task_id": task_id,
                          "parent_task_id": None, "related_task_id": None,
                          "relation_type": "lead" if index == 0 else "collaborate",
                          "person_id": person_id, "role": "lead" if index == 0 else "collaborator",
                          "evidence_id": f"snapshot-report-{person_id}"})
    snapshots = [{"snapshot_id": f"snapshot-report-{person_id}", "source_system": "test-fixture",
                  "source_record_id": f"report-{person_id}", "raw_record": "测试日报工作事项",
                  "evidence_quality": "source-record-id"} for person_id in submitted_ids]
    people_shape = lambda name, index: {"employeeId": index, "name": name, "teamName": "行一二部", "submittedAt": None}
    documents: dict[str, object] = {
        "roster.json": roster,
        "attendance.json": attendance,
        "reports.json": reports,
        "tasks.json": tasks,
        "task-relations.json": relations,
        "projects.json": [{"project_id": "project-main", "standard_name": "重点交付项目",
                           "candidate_names": ["重点交付项目"], "merge_status": "merged",
                           "manual_confirmation_status": "confirmed"}],
        "source-snapshot.json": snapshots,
        "statistics.json": [
            {"snapshot_type": "PROGRESS_1730", "snapshot_date": REPORT_DATE,
             "captured_at": f"{REPORT_DATE}T17:30:00+08:00", "expected_count": 90,
             "submitted_count": 88, "missing_count": 2, "submission_rate": 88 / 90,
             "submitted_people": [], "missing_people": [people_shape(name, index) for index, name in enumerate(MISSING, 1)],
             "late_submitted_people": [], "team_statistics": []},
            {"snapshot_type": "FINAL", "snapshot_date": REPORT_DATE,
             "captured_at": f"{REPORT_DATE}T22:00:00+08:00", "expected_count": 90,
             "submitted_count": 88, "missing_count": 2, "submission_rate": 88 / 90,
             "submitted_people": [], "missing_people": [people_shape(name, index) for index, name in enumerate(MISSING, 1)],
             "late_submitted_people": [], "team_statistics": []},
        ],
        "project-status.json": [
            {"project_id": "project-main", "project_name": "重点交付项目", "snapshot_date": REPORT_DATE,
             "snapshot_captured_at": f"{REPORT_DATE}T22:00:00+08:00", "snapshot_origin": "captured",
             "state": "DELIVERY_IN_PROGRESS", "lifecycle": "delivery", "current_stage": "实施",
             "state_started_date": "2026-07-01", "owner_name": "测试员工001", "participant_count": 88,
             "blocked_task_count": 12, "latest_report_date": REPORT_DATE, "formal": True, "active": True,
             "source_status": "active", "priority": "high"},
            {"project_id": "project-stale-5", "project_name": "存量项目甲", "snapshot_date": REPORT_DATE,
             "snapshot_captured_at": f"{REPORT_DATE}T22:00:00+08:00", "snapshot_origin": "captured",
             "state": "DELIVERY_IN_PROGRESS", "lifecycle": "delivery", "current_stage": "实施",
             "state_started_date": "2026-07-01", "owner_name": "负责人甲", "participant_count": 2,
             "blocked_task_count": 0, "latest_report_date": "2026-07-24", "formal": True, "active": True,
             "source_status": "active", "priority": "high"},
            {"project_id": "project-stale-4", "project_name": "存量项目乙", "snapshot_date": REPORT_DATE,
             "snapshot_captured_at": f"{REPORT_DATE}T22:00:00+08:00", "snapshot_origin": "captured",
             "state": "PRESALES_IN_PROGRESS", "lifecycle": "presales", "current_stage": "方案",
             "state_started_date": "2026-07-01", "owner_name": "负责人乙", "participant_count": 2,
             "blocked_task_count": 0, "latest_report_date": "2026-07-27", "formal": True, "active": True,
             "source_status": "active", "priority": "medium"},
        ],
        "workdays.json": [{"date": day, "workday": True, "source": "DEFAULT", "note": None}
                          for day in ("2026-07-24", "2026-07-27", "2026-07-28", "2026-07-29", "2026-07-30", "2026-07-31")],
    }
    documents["manifest.json"] = {
        "data_contract_version": "1.1.0", "package_id": "leadership-2026-07-31-test-fixture",
        "department_id": "dept-test", "department_name": "行一二部（测试）", "period_type": "daily",
        "start_date": REPORT_DATE, "end_date": REPORT_DATE, "working_dates": [REPORT_DATE],
        "timezone": "Asia/Shanghai", "generated_at": "2026-08-14T10:00:00+08:00",
        "stale_project_threshold_workdays": 3, "file_sha256": {},
    }
    documents["manifest.json"]["file_sha256"] = {
        name: hashlib.sha256(canonical(value)).hexdigest()
        for name, value in documents.items() if name != "manifest.json"
    }
    return documents


def write_package(path: Path) -> Path:
    documents = leadership_documents()
    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name in sorted(documents):
            info = zipfile.ZipInfo(name, date_time=(2026, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(info, canonical(documents[name]))
    return path


def leadership_candidate(facts: dict[str, object]) -> dict[str, object]:
    evidence_id = facts["evidence"][0]["evidence_id"]
    return {
        "data_contract_version": "1.1.0",
        "overall_judgment": {"summary": "当日填报总体稳定，项目推进证据完整。", "person_ids": [],
                             "project_id": None, "evidence_ids": [evidence_id], "limitation_note": None},
        "efficiency_insights": [],
        "project_highlights": [{"summary": "重点交付项目持续推进。", "person_ids": ["p008"],
                                "project_id": "project-main", "evidence_ids": [evidence_id],
                                "limitation_note": None}],
        "risk_items": [],
        "next_day_actions": [{"summary": "跟进阻塞事项并明确反馈时点。", "person_ids": ["p008"],
                              "project_id": "project-main", "evidence_ids": [evidence_id],
                              "limitation_note": None}],
    }


def write_golden_artifacts(package_path: Path, facts_path: Path, report_path: Path) -> None:
    package_result = validate_package(write_package(package_path))
    if not package_result.ok:
        raise ValueError(package_result.errors)
    facts = build_daily_facts(package_result.package, date.fromisoformat(REPORT_DATE))
    candidate = leadership_candidate(facts)
    validation = validate_analysis_result(facts, candidate)
    if not validation.ok:
        raise ValueError(validation.errors)
    facts_path.parent.mkdir(parents=True, exist_ok=True)
    facts_path.write_bytes(canonical(facts))
    render_daily_report(TEMPLATE_SKILL / "assets" / "daily-report-template.docx",
                        facts, candidate, report_path)
    audit = audit_daily_report(report_path, facts)
    if not audit.ok:
        raise ValueError(audit.errors)


class DailyLeadershipPipelineTests(unittest.TestCase):
    def test_data_package_to_validated_leadership_docx(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            package_result = validate_package(write_package(root / "package.zip"))
            self.assertTrue(package_result.ok, package_result.errors)
            facts = build_daily_facts(package_result.package, date.fromisoformat(REPORT_DATE))
            candidate = leadership_candidate(facts)
            validation = validate_analysis_result(facts, candidate)
            self.assertTrue(validation.ok, validation.errors)
            output = render_daily_report(TEMPLATE_SKILL / "assets" / "daily-report-template.docx",
                                         facts, candidate, root / "report.docx")
            audit = audit_daily_report(output, facts)

            self.assertTrue(audit.ok, audit.errors)
            document = Document(output)
            self.assertEqual(len(document.tables), 1)
            self.assertLess(len(document.paragraphs), 60)
            self.assertAlmostEqual(document.sections[0].page_width.inches, 8.5, places=2)
            self.assertEqual(document.paragraphs[0].runs[0].font.name, "黑体")
            self.assertEqual(round(document.paragraphs[0].runs[0].font.size.pt), 18)
            dashboard = facts["attendance_summary"]
            self.assertEqual((dashboard["expected_people"], dashboard["submitted_by_1730_people"],
                              dashboard["submitted_by_2200_people"], dashboard["leave_people"],
                              dashboard["missing_people_count"]), (90, 88, 88, 5, 2))
            self.assertEqual(facts["metrics"]["effective_task_count"], 115)
            self.assertEqual((facts["status_distribution"]["completed"], facts["status_distribution"]["blocked"]), (11, 12))
            self.assertEqual([item["inactive_workdays"] for item in facts["stale_project_alerts"]], [5, 4])
            self.assertEqual([item["name"] for item in dashboard["missing_people"]], list(MISSING))


if __name__ == "__main__":
    unittest.main()
