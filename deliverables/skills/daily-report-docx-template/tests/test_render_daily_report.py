from __future__ import annotations

import tempfile
import unittest
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document


SKILL_ROOT = Path(__file__).resolve().parents[1]
import sys
sys.path.insert(0, str(SKILL_ROOT / "scripts"))

from audit_daily_report import audit_daily_report  # noqa: E402
from render_daily_report import _clean_narrative, render_daily_report  # noqa: E402


TEMPLATE = SKILL_ROOT / "assets" / "daily-report-template.docx"
PEOPLE = lambda *names: [{"person_id": f"p-{index}", "name": name, "work_period_status": None,
                          "review_reason": None, "source_record_id": f"r-{index}"}
                         for index, name in enumerate(names)]
FACTS = {
    "period": {"report_date": "2026-07-31", "start_date": "2026-07-27", "end_date": "2026-07-31", "timezone": "Asia/Shanghai"},
    "metrics": {"expected_people": 90, "submitted_people": 88, "submission_rate": 0.9778,
                "missing_people": 2, "full_day_leave_people": 5, "half_day_leave_people": 2,
                "review_required_people": 0, "effective_task_count": 115, "risk_or_blocked_count": 12},
    "status_distribution": {"completed": 11, "blocked": 12, "in_progress": 92, "planned": 0},
    "attendance_summary": {"expected_people": 90, "submitted_by_1730_people": 88,
                           "submitted_by_1730_rate": 0.9778, "submitted_by_2200_people": 88,
                           "submitted_by_2200_rate": 0.9778, "missing_people": PEOPLE("王达伟", "吴鹏"),
                           "full_day_leave_people": PEOPLE("周菁", "陈苏", "刘茜", "杨瑞", "王海娜"),
                           "half_day_leave_people": [], "review_required_people": [],
                           "calculation_source": "stored_statistics", "leave_people": 5,
                           "missing_people_count": 2},
    "efficiency_summary": [
        {"person_id": "p1", "name": "钱程", "group": "行业解决方案组", "task_count": 8,
         "completed_task_count": 5, "in_progress_task_count": 3, "blocked_task_count": 0,
         "project_count": 2, "lead_task_count": 3, "collaboration_task_count": 1,
         "evidence_ids": ["snapshot-1"]},
        {"person_id": "p2", "name": "丁德胜", "group": "行业解决方案组", "task_count": 6,
         "completed_task_count": 3, "in_progress_task_count": 2, "blocked_task_count": 1,
         "project_count": 3, "lead_task_count": 1, "collaboration_task_count": 2,
         "evidence_ids": ["snapshot-2"]},
    ],
    "formal_project_dynamics": [{"project_id": "project-1", "project_name": "知识库建设项目",
        "state": "DELIVERY_IN_PROGRESS", "current_stage": "实施", "owner_name": "钱程",
        "lead_people": ["钱程"], "collaborator_people": ["丁德胜"], "task_count": 5,
        "completed_task_count": 3, "blocked_task_count": 1, "outputs": ["完成Embedding模型对比"],
        "blockers": ["客户反馈待确认"], "evidence_ids": ["snapshot-1"],
        "snapshot_origin": "captured", "snapshot_captured_at": "2026-07-31T22:00:00+08:00"}],
    "unlinked_project_dynamics": [{"project_id": "candidate-1", "project_name": "连云港健康小屋",
        "state": None, "current_stage": None, "owner_name": None, "lead_people": ["丁德胜"],
        "collaborator_people": [], "task_count": 2, "completed_task_count": 0,
        "blocked_task_count": 0, "outputs": [], "blockers": [], "evidence_ids": ["snapshot-2"],
        "snapshot_origin": None, "snapshot_captured_at": None}],
    "stale_project_alerts": [{"project_id": "project-2", "project_name": "存量交付项目",
        "state": "DELIVERY_IN_PROGRESS", "owner_name": "王经理", "latest_report_date": "2026-07-24",
        "current_stage": "实施", "inactive_workdays": 4, "threshold_workdays": 3,
        "blocked_task_count": 0, "snapshot_origin": "captured", "snapshot_captured_at": None}],
    "project_continuity": [{"project_id": "project-1", "project_name": "知识库建设项目",
        "active_dates": ["2026-07-27", "2026-07-29", "2026-07-31"],
        "first_active_date": "2026-07-27", "last_active_date": "2026-07-31",
        "person_ids": ["p1", "p2"], "people": ["钱程", "丁德胜"],
        "actions": ["模型选型", "技术方案", "部署测试"], "outputs": ["完成Embedding模型对比"],
        "status_evidence": ["completed", "in_progress"], "classification": "continuing",
        "limitation_note": None, "evidence_ids": ["snapshot-1", "snapshot-2"]}],
    "project_associations": [{"association_id": "project-1-multi-person",
        "association_type": "same_project_multi_person", "project_ids": ["project-1"],
        "person_ids": ["p1", "p2"], "people": ["钱程", "丁德胜"],
        "description": "知识库建设项目由钱程主导、丁德胜协同推进。",
        "evidence_ids": ["snapshot-1", "snapshot-2"]}],
    "risk_assessment": [{"risk_id": "project-progress-project-1", "risk_type": "project_progress",
        "level": "medium", "project_id": "project-1", "person_ids": ["p2"],
        "summary": "知识库建设项目存在客户反馈待确认事项。", "evidence_ids": ["snapshot-2"],
        "limitation_note": None}],
    "risk_candidates": [{"task_id": "task-1", "person_id": "p2", "status": "blocked",
                         "risk": "客户反馈待确认", "blocker": None}],
    "data_quality": {"project_status_coverage": "available", "reconstructed_project_count": 0,
                     "limitations": []},
}
ANALYSIS = {
    "data_contract_version": "1.2.0",
    "overall_judgment": {"summary": "当日工作以项目交付和方案支撑为主，填报总体及时。", "person_ids": [],
                         "project_id": None, "evidence_ids": ["snapshot-1"], "limitation_note": None},
    "efficiency_insights": [{"summary": "钱程（employee-27）主导知识库建设项目（project-1），完成Embedding模型对比，并推进部署测试。",
                              "person_ids": ["p1"], "project_id": "project-1",
                              "evidence_ids": ["snapshot-1"], "limitation_note": None}],
    "continuity_analysis": [{"project_id": "project-1", "summary": "知识库建设项目在三个工作日持续推进，已从模型选型进入部署测试。",
                              "person_ids": ["p1", "p2"], "evidence_ids": ["snapshot-1", "snapshot-2"], "limitation_note": None}],
    "association_analysis": [{"project_id": "project-1", "summary": "钱程主导模型验证，丁德胜协同准备材料，分工关系明确。",
                               "person_ids": ["p1", "p2"], "evidence_ids": ["snapshot-1", "snapshot-2"], "limitation_note": None}],
    "risk_items": [{"summary": "客户反馈尚待确认。", "person_ids": [], "project_id": None,
                    "evidence_ids": ["snapshot-2"], "limitation_note": None}],
    "next_day_actions": [{"summary": "跟进客户反馈并明确责任人和完成时点。", "person_ids": [], "project_id": None,
                          "evidence_ids": ["snapshot-2"], "limitation_note": None}],
}


class RenderDailyReportTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tempdir = tempfile.TemporaryDirectory()
        self.output = Path(self.tempdir.name) / "daily-report.docx"

    def tearDown(self) -> None:
        self.tempdir.cleanup()

    def test_facts_only_report_is_complete_and_leader_readable(self) -> None:
        path = render_daily_report(TEMPLATE, FACTS, None, self.output)
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs +
                         [p for table in document.tables for row in table.rows for cell in row.cells for p in cell.paragraphs])
        audit = audit_daily_report(path, FACTS)

        self.assertTrue(audit.ok, audit.errors)
        for required in ("行一二部工作日报", "2026年7月31日",
                         "截至17:30", "截至22:00", "王达伟", "吴鹏", "周菁",
                         "人员效能分析", "钱程", "知识库建设项目", "连云港健康小屋",
                         "存量交付项目", "当日项目动态", "项目关联性及协同分析",
                         "风险评估", "管理建议"):
            self.assertIn(required, text)
        self.assertNotIn("待人工补充", text)
        self.assertNotIn("领导版", text)
        self.assertNotIn("数据时区", text)
        self.assertNotIn("PRESALES_IN_PROGRESS", text)
        for metric in ("90", "88", "5", "2", "115", "11", "12"):
            self.assertIn(metric, text)
        self.assertEqual(len(document.tables[0].rows), 2)
        self.assertEqual(len(document.tables[0].columns), 4)
        self.assertEqual(len(document.tables), 1)
        headings = [paragraph.text for paragraph in document.paragraphs
                    if paragraph.style.name in {"Heading 1", "Heading 2"}]
        self.assertEqual(headings, [
            "一、总体概况", "二、填报与出勤情况", "三、人员效能分析",
            "四、当日项目动态", "五、项目关联性及协同分析",
            "六、风险评估", "七、管理建议",
        ])
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        self.assertGreaterEqual(xml.count('w:fill="1F4E78"'), 8)
        self.assertIn("微软雅黑", xml)
        self.assertIn("宋体", xml)

    def test_valid_analysis_enriches_without_replacing_deterministic_sections(self) -> None:
        path = render_daily_report(TEMPLATE, FACTS, ANALYSIS, self.output)
        document = Document(path)
        text = "\n".join(p.text for p in document.paragraphs +
                         [p for table in document.tables for row in table.rows for cell in row.cells for p in cell.paragraphs])

        self.assertIn("当日工作以项目交付和方案支撑为主", text)
        self.assertIn("钱程主导知识库建设项目，完成Embedding模型对比，并推进部署测试", text)
        self.assertIn("知识库建设项目在三个工作日持续推进", text)
        self.assertIn("钱程主导模型验证", text)
        self.assertIn("王达伟", text)
        self.assertNotIn("employee-27", text)
        self.assertNotIn("project-1", text)
        self.assertNotIn("（无）", text)

    def test_ai_continuity_supplements_formal_projects_not_selected_by_model(self) -> None:
        facts = deepcopy(FACTS)
        omitted_project = deepcopy(facts["formal_project_dynamics"][0])
        omitted_project.update({"project_id": "project-99", "project_name": "补充展示项目"})
        facts["formal_project_dynamics"].append(omitted_project)
        facts["project_continuity"].append({
            "project_id": "project-99", "project_name": "补充展示项目",
            "active_dates": ["2026-07-31"], "people": ["钱程"],
            "actions": ["方案编制"], "outputs": ["形成初稿"], "classification": "new",
        })

        path = render_daily_report(TEMPLATE, facts, ANALYSIS, self.output)
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("补充展示项目", text)
        self.assertTrue(audit_daily_report(path, facts).ok)

    def test_internal_placeholders_and_data_quality_notes_are_hidden(self) -> None:
        facts = deepcopy(FACTS)
        facts["data_quality"] = {"project_status_coverage": "available", "reconstructed_project_count": 1,
                                 "limitations": ["存量交付项目的历史状态为事后重建（采集时间：2026-08-14T10:00:00+08:00）"]}
        facts["project_continuity"].append({
            "project_id": "project-3", "project_name": "内部专项",
            "active_dates": ["2026-07-31"], "first_active_date": "2026-07-31",
            "last_active_date": "2026-07-31", "person_ids": ["p1"], "people": ["钱程"],
            "actions": ["内部管理"], "outputs": [], "status_evidence": ["in_progress"],
            "classification": "new", "limitation_note": "仅为内部占位分类。",
            "evidence_ids": ["snapshot-1"],
        })
        facts["unlinked_project_dynamics"].extend([
            {**deepcopy(facts["unlinked_project_dynamics"][0]), "project_id": "project-2",
             "project_name": "暂无正式项目"},
            {**deepcopy(facts["unlinked_project_dynamics"][0]), "project_id": "project-3",
             "project_name": "内部专项"},
        ])

        path = render_daily_report(TEMPLATE, facts, None, self.output)
        document = Document(path)
        text = "\n".join(p.text for p in document.paragraphs)
        audit = audit_daily_report(path, facts)

        self.assertTrue(audit.ok, audit.errors)
        self.assertNotIn("暂无正式项目", text)
        self.assertNotIn("内部专项", text)
        self.assertNotIn("数据口径说明", text)
        self.assertNotIn("事后重建", text)

    def test_facts_only_report_uses_project_metrics_not_raw_daily_text(self) -> None:
        facts = deepcopy(FACTS)
        raw_daily_text = "完成省应急厅项目招标文件编制，随后开展组内审核、客户交叉审核，并根据审核意见优化调整文稿后交付客户。"
        facts["project_continuity"][0]["actions"] = [raw_daily_text]
        facts["project_continuity"][0]["outputs"] = [raw_daily_text]
        facts["formal_project_dynamics"][0]["outputs"] = [raw_daily_text]
        path = render_daily_report(TEMPLATE, facts, None, self.output)
        document = Document(path)
        text = "\n".join(p.text for p in document.paragraphs)

        self.assertIn("钱程：", text)
        self.assertIn("参与知识库建设项目", text)
        self.assertIn("当日记录8项事项", text)
        self.assertIn("详情见原始日报证据", text)
        self.assertNotIn(raw_daily_text, text)

    def test_fallback_formal_project_summary_does_not_render_raw_output(self) -> None:
        facts = deepcopy(FACTS)
        raw_daily_text = "完成省应急厅项目招标文件编制，先后开展组内审核和客户交叉审核，根据意见优化后交付客户。"
        facts["project_continuity"] = []
        facts["formal_project_dynamics"][0]["outputs"] = [raw_daily_text]
        path = render_daily_report(TEMPLATE, facts, None, self.output)
        document = Document(path)
        text = "\n".join(p.text for p in document.paragraphs)

        self.assertIn("知识库建设项目，负责人钱程", text)
        self.assertIn("当日记录5项事项，其中完成3项、阻塞1项", text)
        self.assertIn("详情见原始日报证据", text)
        self.assertNotIn(raw_daily_text, text)

    def test_readability_cleanup_preserves_business_parentheses(self) -> None:
        cleaned = _clean_narrative("钱程（employee-27）推进采购项目（标段二）（project-1），当前停滞4天。")

        self.assertEqual(cleaned, "钱程推进采购项目（标段二），当前停滞4天。")

    def test_readability_cleanup_translates_project_state_codes(self) -> None:
        cleaned = _clean_narrative(
            "项目历史状态为事后重建(PRESALES_IN_PROGRESS)，交付状态为DELIVERY_IN_PROGRESS，"
            "售后状态为AFTERSALES_IN_PROGRESS/operations-support，阻塞状态为BLOCKED；状态证据为blocked和completed。"
        )

        self.assertEqual(
            cleaned,
            "项目历史状态为事后重建（售前推进中），交付状态为交付推进中，售后状态为售后运维中/运维保障，阻塞状态为受阻；状态证据为受阻和已完成。",
        )

    def test_readability_cleanup_replaces_imported_placeholder_and_unknown_status_code(self) -> None:
        cleaned = _clean_narrative(
            "阻塞说明统一标记为 Imported blocker needs external coordination，当前状态为 FUTURE_EXTERNAL_STATUS。"
        )

        self.assertEqual(cleaned, "阻塞说明统一标记为存在阻塞事项，待补充具体原因，当前状态为未识别状态。")

    def test_readability_cleanup_removes_internal_project_classification_clause(self) -> None:
        cleaned = _clean_narrative(
            "正式在管项目共4个，其余动态均归入内部专项或暂无正式项目，跨组协同以疾控AI项目为主。"
        )

        self.assertEqual(cleaned, "正式在管项目共4个，跨组协同以疾控AI项目为主。")

    def test_readability_cleanup_removes_internal_task_identifiers_in_parentheses(self) -> None:
        cleaned = _clean_narrative("明确记录为阻塞状态（task-5703-afternoon/morning），请跟进客户侧验证。")

        self.assertEqual(cleaned, "明确记录为阻塞状态，请跟进客户侧验证。")

    def test_readability_cleanup_hides_internal_project_management_fields(self) -> None:
        cleaned = _clean_narrative(
            "项目由吴海宁主导（lead），分类为new。"
            "在留存的候选项目身份（merge_status=unmerged, manual_confirmation_status=pending）下，"
            "仍未与正式项目主数据完成合并；补齐项目主数据：落实项目的merge_status与"
            "manual_confirmation_status，正式确认项目owner。"
            "数据包未提供项目状态快照（reconstructed_project_count=0, "
            "project_status_coverage=unavailable）。"
        )

        self.assertEqual(
            cleaned,
            "项目由吴海宁主导，分类为新增动态。数据包未提供项目状态快照。",
        )

    def test_unlinked_project_is_reported_without_manual_association_prompt(self) -> None:
        path = render_daily_report(TEMPLATE, FACTS, None, self.output)
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("日报中新出现项目", text)
        self.assertIn("已纳入当日项目动态统计", text)
        self.assertNotIn("确认关联关系", text)
        self.assertNotIn("项目维护主数据", text)

    def test_project_already_described_by_ai_is_not_rendered_twice(self) -> None:
        analysis = deepcopy(ANALYSIS)
        analysis["continuity_analysis"] = [{
            "project_id": None,
            "summary": "连云港健康小屋当日出现项目动态，由丁德胜推进需求调研。",
            "person_ids": ["p2"],
            "evidence_ids": ["snapshot-2"],
            "limitation_note": None,
        }]
        path = render_daily_report(TEMPLATE, FACTS, analysis, self.output)
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertEqual(1, text.count("连云港健康小屋"))
        self.assertIn("该项目为日报中新出现项目，即首次出现在本期日报记录中", text)
        self.assertNotIn("日报中新出现项目：", text)

    def test_valid_analysis_avoids_verbose_risk_duplicates_and_keeps_stale_alerts(self) -> None:
        facts = deepcopy(FACTS)
        facts["risk_assessment"].extend([
            {"summary": "暂无正式项目存在明确阻塞或风险记录。"},
            {"summary": "内部专项存在明确阻塞或风险记录。"},
        ])

        path = render_daily_report(TEMPLATE, facts, ANALYSIS, self.output)
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("客户反馈尚待确认", text)
        self.assertIn("存量交付项目已连续4个工作日无日报动态", text)
        self.assertEqual(1, text.count("存量交付项目已连续4个工作日无日报动态"))
        self.assertNotIn("暂无正式项目", text)
        self.assertNotIn("内部专项", text)


if __name__ == "__main__":
    unittest.main()
