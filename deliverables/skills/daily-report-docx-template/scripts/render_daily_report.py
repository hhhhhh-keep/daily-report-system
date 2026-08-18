"""Deterministically render a concise leader-facing daily report DOCX."""

from __future__ import annotations

import argparse
import re
from collections.abc import Mapping
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ANCHORS = ("DR_TITLE", "DR_META", "DR_DASHBOARD", "DR_ATTENDANCE", "DR_WORK_SUMMARY",
           "DR_PROJECTS", "DR_RISKS", "DR_ACTIONS")
BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
HIDDEN_PROJECT_NAMES = {"暂无正式项目", "内部专项"}
TECHNICAL_ID = re.compile(r"(?:employee|project|task)-[A-Za-z0-9_/-]+", re.IGNORECASE)
TECHNICAL_PAREN = re.compile(r"[（(]([^（）()]*(?:employee|project|task)-[^（）()]*)[）)]", re.IGNORECASE)
IMPORTED_BLOCKER = "Imported blocker needs external coordination"
UNKNOWN_STATUS_CODE = re.compile(r"\b[A-Z][A-Z0-9]*(?:_[A-Z0-9]+)+\b")


def create_template(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document)
    for anchor in ANCHORS:
        document.add_paragraph(anchor)
    document.save(path)
    return path


def render_daily_report(template_path: Path, facts: Mapping[str, object],
                        analysis: Mapping[str, object] | None, output_path: Path) -> Path:
    document = Document(template_path)
    _clear_body(document)
    period = _mapping(facts.get("period"))
    attendance = _mapping(facts.get("attendance_summary"))
    metrics = _mapping(facts.get("metrics"))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    title.paragraph_format.first_line_indent = Pt(0)
    _add_run(title, "行一二部工作日报", "黑体", 18, bold=True)
    meta = document.add_paragraph()
    meta.paragraph_format.line_spacing = 1.5
    report_date = str(period.get("report_date") or period.get("end_date") or period.get("start_date") or "")
    _add_run(meta, _chinese_date(report_date), "仿宋_GB2312", 11)

    _dashboard(document, attendance, metrics, _mapping(facts.get("status_distribution")))
    _heading(document, "一、总体概况")
    _body(document, _overall_judgment(attendance, metrics, analysis))

    _heading(document, "二、填报与出勤情况")
    _attendance_narrative(document, attendance)

    _heading(document, "三、人员效能分析")
    _efficiency_narrative(document, facts, analysis)

    _heading(document, "四、当日项目动态")
    _continuity_narrative(document, facts, analysis)

    _heading(document, "五、项目关联性及协同分析")
    _association_narrative(document, facts, analysis)

    _heading(document, "六、风险评估")
    _risk_narrative(document, facts, analysis)
    _heading(document, "七、管理建议")
    _action_narrative(document, facts, analysis)
    if analysis is None:
        _body(document, "注：本报告未采用AI语义补充，以上内容均由确定性规则和原始证据生成。", italic=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.5), Inches(11.0)
    section.top_margin = section.bottom_margin = Inches(0.83)
    section.left_margin = section.right_margin = Inches(0.83)
    normal = document.styles["Normal"]
    _set_style_font(normal, "仿宋_GB2312", 11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(23.75)
    heading1 = document.styles["Heading 1"]
    _set_style_font(heading1, "黑体", 15)
    heading1.paragraph_format.line_spacing = 1.5
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(0)
    heading1.paragraph_format.first_line_indent = Pt(0)
    heading1.paragraph_format.keep_with_next = True
    heading2 = document.styles["Heading 2"]
    _set_style_font(heading2, "黑体", 13)
    heading2.paragraph_format.line_spacing = 1.5
    heading2.paragraph_format.space_before = Pt(0)
    heading2.paragraph_format.space_after = Pt(0)
    heading2.paragraph_format.first_line_indent = Pt(0)
    heading2.paragraph_format.keep_with_next = True


def _dashboard(document: Document, attendance: Mapping[str, object], metrics: Mapping[str, object],
               status: Mapping[str, object]) -> None:
    labels = ("应填写", "17:30已填", "22:00已填", "请假", "未填写", "有效事项", "完成事项", "阻塞事项")
    leave_count = len(_list(attendance.get("full_day_leave_people"))) + len(_list(attendance.get("half_day_leave_people")))
    values = (attendance.get("expected_people"), attendance.get("submitted_by_1730_people"),
              attendance.get("submitted_by_2200_people"), leave_count,
              len(_list(attendance.get("missing_people"))), metrics.get("effective_task_count", 0),
              status.get("completed", 0), status.get("blocked", metrics.get("risk_or_blocked_count", 0)))
    table = document.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (label, value) in enumerate(zip(labels, values)):
        cell = table.cell(index // 4, index % 4)
        cell.width = Cm(4.0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        _add_run(paragraph, label + "\n", "微软雅黑", 9, color="FFFFFF")
        _add_run(paragraph, _display(value), "宋体", 16, color="FFFFFF", bold=True)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _overall_judgment(attendance: Mapping[str, object], metrics: Mapping[str, object],
                      analysis: Mapping[str, object] | None) -> str:
    if analysis is not None:
        summary = _mapping(analysis.get("overall_judgment")).get("summary")
        if not isinstance(summary, str):
            summary = _mapping(analysis.get("work_summary")).get("summary")
        if isinstance(summary, str) and summary.strip() and not _is_placeholder(summary):
            return summary.strip()
    expected = attendance.get("expected_people")
    final_rate = _percent(attendance.get("submitted_by_2200_rate"))
    missing = len(_list(attendance.get("missing_people")))
    risks = metrics.get("risk_or_blocked_count", 0)
    tasks = metrics.get("effective_task_count", 0)
    submitted = metrics.get("submitted_people", 0)
    return (f"当日应填写{_display(expected)}人，实际有效填写{_display(submitted)}人，截至22:00填报率"
            f"{final_rate}，未填写{missing}人；共形成{_display(tasks)}项有效事项，其中风险或阻塞事项"
            f"{_display(risks)}项。管理上需同步关注填报完整性、项目连续推进情况、人员协同关系及明确风险。")


def _continuity_narrative(document: Document, facts: Mapping[str, object],
                          analysis: Mapping[str, object] | None) -> None:
    ai_items = [item for item in _list(analysis.get("continuity_analysis") if analysis else None)
                if not _contains_hidden_project(str(item.get("summary") or ""))]
    deterministic = [item for item in _list(facts.get("project_continuity"))
                     if _is_visible_project(item.get("project_name"))]
    rendered = False
    if ai_items:
        for index, item in enumerate(ai_items, 1):
            _numbered_body(document, index, str(item.get("summary")))
        rendered = True
    elif deterministic:
        classification_names = {
            "new": "本周期新增动态", "continuing": "本周期持续推进",
            "interrupted": "后段未见连续动态", "last_activity_completed": "末次事项已有完成记录",
        }
        for index, item in enumerate(deterministic, 1):
            active_dates = [_chinese_date(value, include_year=False) for value in _string_list(item.get("active_dates"))]
            task_count = _project_task_count(facts, item)
            _numbered_body(
                document, index,
                f"{_display(item.get('project_name'))}：{classification_names.get(item.get('classification'), '窗口内有动态')}；"
                f"动态日期为{_join(active_dates)}；参与人员{_join(item.get('people'))}；"
                f"当日记录{task_count}项事项，已形成项目推进动态，详情见原始日报证据。",
            )
        rendered = True

    known = ({item.get("project_id") for item in ai_items}
             if ai_items else {item.get("project_id") for item in deterministic})
    supplements = [item for item in _list(facts.get("formal_project_dynamics"))
                   if item.get("project_id") not in known and _is_visible_project(item.get("project_name"))]
    if supplements:
        _body(document, "当日项目补充：")
        _formal_projects(document, supplements, None)
        rendered = True
    unlinked = [item for item in _list(facts.get("unlinked_project_dynamics"))
                if _is_visible_project(item.get("project_name"))]
    if unlinked:
        _body(document, "日报识别待关联项目：")
        _unlinked_projects(document, unlinked)
        rendered = True
    if not rendered:
        _body(document, "当日未识别到项目动态。")


def _association_narrative(document: Document, facts: Mapping[str, object],
                           analysis: Mapping[str, object] | None) -> None:
    ai_items = _list(analysis.get("association_analysis") if analysis else None)
    if ai_items:
        for index, item in enumerate(ai_items, 1):
            _numbered_body(document, index, str(item.get("summary")))
        return
    associations = _list(facts.get("project_associations"))
    if associations:
        for index, item in enumerate(associations, 1):
            _numbered_body(document, index, str(item.get("description")))
        return
    derived = []
    for item in _list(facts.get("formal_project_dynamics")):
        people = sorted(set((*_string_list(item.get("lead_people")), *_string_list(item.get("collaborator_people")))))
        if people:
            derived.append(f"{item.get('project_name')}由{_join(item.get('lead_people'))}主导、{_join(item.get('collaborator_people'))}协同，涉及人员{_join(people)}。")
    if derived:
        for index, item in enumerate(derived, 1):
            _numbered_body(document, index, item)
    else:
        _body(document, "当日未形成可核验的多人或跨项目协同关系。")


def _attendance_narrative(document: Document, attendance: Mapping[str, object]) -> None:
    items = (
        f"应填写人数：{_display(attendance.get('expected_people'))}人。",
        f"截至17:30：已填写{_display(attendance.get('submitted_by_1730_people'))}人，填写率{_percent(attendance.get('submitted_by_1730_rate'))}。",
        f"截至22:00：已填写{_display(attendance.get('submitted_by_2200_people'))}人，填写率{_percent(attendance.get('submitted_by_2200_rate'))}。",
        f"未填写人员：{_names(attendance.get('missing_people'))}。",
        f"全天请假人员：{_names(attendance.get('full_day_leave_people'))}。",
        f"半天请假人员：{_names(attendance.get('half_day_leave_people'))}。",
        f"待核验人员：{_names(attendance.get('review_required_people'))}。",
    )
    for index, item in enumerate(items, 1):
        _numbered_body(document, index, item)


def _efficiency_narrative(document: Document, facts: Mapping[str, object],
                          analysis: Mapping[str, object] | None) -> None:
    items = _list(facts.get("efficiency_summary"))
    if not items:
        _body(document, "当日无可统计的人员工作事项。")
        return

    total_tasks = sum(int(item.get("task_count") or 0) for item in items)
    total_completed = sum(int(item.get("completed_task_count") or 0) for item in items)
    total_in_progress = sum(int(item.get("in_progress_task_count") or 0) for item in items)
    total_blocked = sum(int(item.get("blocked_task_count") or 0) for item in items)
    _body(
        document,
        f"当日纳入{len(items)}名人员、{total_tasks}项有效事项，其中完成{total_completed}项、"
        f"进行中{total_in_progress}项、阻塞{total_blocked}项。以下结合具体项目、动作、产出和协同情况进行观察，"
        "不以事项数量直接评价绩效。",
    )

    ai_items = [item for item in _list(analysis.get("efficiency_insights") if analysis else None)
                if str(item.get("summary") or "").strip()]
    if ai_items:
        for index, item in enumerate(ai_items, 1):
            _numbered_body(document, index, str(item.get("summary")))
        return

    context = _efficiency_context(facts)
    displayed = items[:10]
    for index, item in enumerate(displayed, 1):
        name = _display(item.get("name"))
        details = context.get(name, {})
        projects = _string_list(details.get("projects"))[:3]
        summary = (f"{name}：当日记录{_display(item.get('task_count'))}项事项，完成"
                   f"{_display(item.get('completed_task_count'))}项、进行中"
                   f"{_display(item.get('in_progress_task_count'))}项、阻塞"
                   f"{_display(item.get('blocked_task_count'))}项")
        if projects:
            summary += f"；参与{_join(projects)}"
        if projects:
            summary += "，已记录相关项目推进动态，详情见原始日报证据"
        lead = int(item.get("lead_task_count") or 0)
        collaboration = int(item.get("collaboration_task_count") or 0)
        if lead or collaboration:
            summary += f"；其中主导{lead}项、协同{collaboration}项"
        summary += "。"
        _numbered_body(document, index, summary)


def _efficiency_context(facts: Mapping[str, object]) -> dict[str, dict[str, list[str]]]:
    context: dict[str, dict[str, list[str]]] = {}

    def add(name: object, project: object) -> None:
        if not isinstance(name, str) or not name or not _is_visible_project(project):
            return
        row = context.setdefault(name, {"projects": []})
        text = str(project).strip()
        if text and text not in row["projects"]:
            row["projects"].append(text)

    for item in _list(facts.get("project_continuity")):
        for name in _string_list(item.get("people")):
            add(name, item.get("project_name"))
    for section in ("formal_project_dynamics", "unlinked_project_dynamics"):
        for item in _list(facts.get(section)):
            for name in (*_string_list(item.get("lead_people")), *_string_list(item.get("collaborator_people"))):
                add(name, item.get("project_name"))
    return context


def _formal_projects(document: Document, value: object, analysis: Mapping[str, object] | None) -> None:
    items = _list(value)
    if not items:
        _body(document, "当日无已关联正式项目动态。")
        return
    highlights = {item.get("project_id"): item.get("summary") for item in _list(analysis.get("project_highlights") if analysis else None)}
    for index, item in enumerate(items, 1):
        people = "主导" + _join(item.get("lead_people")) + "，协同" + _join(item.get("collaborator_people"))
        highlight = highlights.get(item.get("project_id"))
        task_count = _display(item.get("task_count"))
        completed = _display(item.get("completed_task_count"))
        blocked = int(item.get("blocked_task_count") or 0)
        dynamic_summary = highlight or (
            f"当日记录{task_count}项事项，其中完成{completed}项、阻塞{blocked}项；"
            "已形成项目推进动态，详情见原始日报证据"
        )
        blocker_summary = (
            "阻塞情况：存在阻塞事项，待补充具体原因。" if blocked > 0 else "阻塞情况：无。"
        )
        summary = (
            f"{_display(item.get('project_name'))}，负责人{item.get('owner_name') or '未维护'}："
            f"当前状态/阶段为{_display(item.get('state'))}/{_display(item.get('current_stage'))}；"
            f"当日{people}；{dynamic_summary}；"
            f"{blocker_summary}"
        )
        _numbered_body(document, index, summary)


def _unlinked_projects(document: Document, value: object) -> None:
    items = _list(value)
    if not items:
        _body(document, "无。")
        return
    for index, item in enumerate(items, 1):
        people = sorted(set((*_string_list(item.get("lead_people")), *_string_list(item.get("collaborator_people")))))
        _numbered_body(
            document, index,
            f"{_display(item.get('project_name'))}：涉及人员{_join(people)}，当日事项"
            f"{_display(item.get('task_count'))}项。该名称来自日报语义识别，需核对项目维护主数据后确认关联关系。",
        )


def _project_task_count(facts: Mapping[str, object], item: Mapping[str, object]) -> str:
    project_id = item.get("project_id")
    for section in ("formal_project_dynamics", "unlinked_project_dynamics"):
        for candidate in _list(facts.get(section)):
            if candidate.get("project_id") == project_id:
                return _display(candidate.get("task_count"))
    return "若干"


def _stale_projects(document: Document, value: object) -> None:
    items = _list(value)
    if not items:
        _body(document, "无。")
        return
    for index, item in enumerate(items, 1):
        _numbered_body(
            document, index,
            f"{_display(item.get('project_name'))}，负责人{item.get('owner_name') or '未维护'}："
            f"最近日报动态为{_display(item.get('latest_report_date'))}，已连续"
            f"{_display(item.get('inactive_workdays'))}个工作日无动态，项目维护状态为"
            f"{_display(item.get('state'))}，建议核实当前进展。",
        )


def _risk_narrative(document: Document, facts: Mapping[str, object],
                    analysis: Mapping[str, object] | None) -> None:
    risks = [str(item.get("summary") or item.get("description"))
             for item in _list(analysis.get("risk_items") if analysis else None)
             if (item.get("summary") or item.get("description"))
             and not _contains_hidden_project(str(item.get("summary") or item.get("description")))]
    if not risks:
        for item in _list(facts.get("risk_assessment")):
            summary = str(item.get("summary") or "")
            if summary and not _contains_hidden_project(summary):
                risks.append(summary)
        for item in _list(facts.get("risk_candidates")):
            risk = str(item.get("risk") or item.get("blocker") or "")
            if risk and not _contains_hidden_project(risk):
                risks.append(risk)
    for item in _list(facts.get("stale_project_alerts")):
        project_name = str(item.get("project_name") or "")
        if (_contains_hidden_project(project_name)
                or any(project_name and project_name in risk for risk in risks)):
            continue
        risks.append(
            f"{project_name}已连续{item.get('inactive_workdays')}个工作日无日报动态，"
            f"项目维护状态为{_display(item.get('state'))}，需核验当前进展。"
        )
    risks = list(dict.fromkeys(risks))
    if not risks:
        _body(document, "当日无明确风险或需升级协调事项。")
        return
    for index, risk in enumerate(risks, 1):
        _numbered_body(document, index, str(risk))


def _action_narrative(document: Document, facts: Mapping[str, object],
                      analysis: Mapping[str, object] | None) -> None:
    actions = [(item.get("summary") or item.get("action"))
               for item in _list(analysis.get("next_day_actions") if analysis else None)
               if item.get("summary") or item.get("action")]
    if not actions:
        attendance = _mapping(facts.get("attendance_summary"))
        missing = _names(attendance.get("missing_people"))
        if missing != "无":
            actions.append(f"督促未填人员补报：{missing}。")
        for item in _list(facts.get("stale_project_alerts")):
            actions.append(f"核查{item.get('project_name')}连续{item.get('inactive_workdays')}个工作日无日报动态的原因。")
    if not actions:
        _body(document, "次工作日按既定计划推进。")
        return
    for index, action in enumerate(dict.fromkeys(actions), 1):
        _numbered_body(document, index, str(action))


def _heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.first_line_indent = Pt(0)
    _add_run(paragraph, text, "黑体", 15)


def _subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    _add_run(paragraph, text, "黑体", 13)


def _body(document: Document, text: str, *, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = _add_run(paragraph, _clean_narrative(text), "仿宋_GB2312", 11)
    run.italic = italic


def _numbered_body(document: Document, index: int, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(21)
    paragraph.paragraph_format.first_line_indent = Pt(-9.9)
    _add_run(paragraph, f"{index}. {_clean_narrative(text)}", "仿宋_GB2312", 11)


def _is_visible_project(value: object) -> bool:
    return isinstance(value, str) and value.strip() not in HIDDEN_PROJECT_NAMES


def _contains_hidden_project(text: str) -> bool:
    return any(name in text for name in HIDDEN_PROJECT_NAMES)


def _clean_narrative(text: object) -> str:
    value = str(text or "")

    value = value.replace(IMPORTED_BLOCKER, "存在阻塞事项，待补充具体原因")

    value = re.sub(
        r"其余动态[^，,。；;]*(?:内部专项|暂无正式项目)[^，,。；;]*[，,]?",
        "",
        value,
    )

    for code, label in (
        ("PRESALES_IN_PROGRESS", "售前推进中"),
        ("DELIVERY_IN_PROGRESS", "交付推进中"),
        ("AFTERSALES_IN_PROGRESS", "售后运维中"),
        ("BLOCKED", "受阻"),
        ("in_progress", "进行中"),
        ("blocked", "受阻"),
        ("completed", "已完成"),
        ("planned", "计划中"),
        ("operations-support", "运维保障"),
    ):
        value = value.replace(code, label)
        value = value.replace(f"({label})", f"（{label}）")

    value = UNKNOWN_STATUS_CODE.sub("未识别状态", value)
    value = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", value)

    def unwrap(match: re.Match[str]) -> str:
        remaining = TECHNICAL_ID.sub("", match.group(1))
        remaining = remaining.strip(" ,，、/;；:：")
        remaining = remaining.replace("/", "、")
        return f"，{remaining}" if remaining else ""

    value = TECHNICAL_PAREN.sub(unwrap, value)
    value = TECHNICAL_ID.sub("", value)
    value = re.sub(r"[（(]\s*(?:无|未维护)?\s*[）)]", "", value)
    value = re.sub(r"\s+([，。；：、])", r"\1", value)
    value = re.sub(r"([，、；]){2,}", r"\1", value)
    return value.strip()


def _key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = _table(document, ("项目", "情况"), (3.2, 12.8))
    for label, value in rows:
        _fill_row(table.add_row().cells, (label, value))


def _table(document: Document, headers: tuple[str, ...], widths: tuple[float, ...]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        cell.width = Cm(width)
        _shade_cell(cell, LIGHT_BLUE)
        _set_cell_text(cell, header, bold=True, center=True)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def _fill_row(cells: object, values: tuple[object, ...]) -> None:
    for cell, value in zip(cells, values):
        _set_cell_text(cell, _display(value), center=not isinstance(value, str) or len(value) < 12)


def _set_cell_text(cell: object, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(18)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    _add_run(paragraph, text, "仿宋_GB2312", 10.5, bold=bold)


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_style_font(style: object, name: str, size: float) -> None:
    style.font.name, style.font.size = name, Pt(size)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def _add_run(paragraph: object, text: str, font: str, size: float, *, color: str | None = None,
             bold: bool = False):
    run = paragraph.add_run(text)
    run.font.name, run.font.size, run.bold = font, Pt(size), bold
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _shade_cell(cell: object, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _mapping(value: object) -> Mapping[str, object]:
    return value if isinstance(value, Mapping) else {}


def _list(value: object) -> list[Mapping[str, object]]:
    return [item for item in value if isinstance(item, Mapping)] if isinstance(value, list) else []


def _names(value: object) -> str:
    return _join([item.get("name") for item in _list(value) if item.get("name")])


def _join(value: object) -> str:
    items = value if isinstance(value, (list, tuple, set)) else []
    texts = [str(item) for item in items if item not in (None, "")]
    return "、".join(texts) if texts else "无"


def _string_list(value: object) -> list[str]:
    return [str(item) for item in value if item not in (None, "")] if isinstance(value, (list, tuple, set)) else []


def _display(value: object) -> str:
    state_names = {
        "DELIVERY_IN_PROGRESS": "交付推进中",
        "PRESALES_IN_PROGRESS": "售前推进中",
        "AFTERSALES_IN_PROGRESS": "售后运维中",
        "BLOCKED": "受阻",
        "in_progress": "进行中",
        "blocked": "受阻",
        "completed": "已完成",
        "planned": "计划中",
        "COMPLETED": "已完成",
        "PAUSED": "暂停",
        "implementation": "实施阶段",
        "solution-design": "方案设计阶段",
        "proposal": "方案阶段",
        "operations-support": "运维保障",
    }
    if value in (None, ""):
        return "无"
    return state_names.get(str(value), str(value))


def _percent(value: object) -> str:
    return "待核验" if value is None else f"{float(value):.2%}"


def _chinese_date(value: object, *, include_year: bool = True) -> str:
    text = str(value or "")
    try:
        year, month, day = (int(part) for part in text[:10].split("-"))
    except (TypeError, ValueError):
        return text or "未提供"
    return f"{year}年{month}月{day}日" if include_year else f"{month}月{day}日"


def _is_placeholder(value: str) -> bool:
    normalized = " ".join(value.lower().split())
    return any(phrase in normalized for phrase in (
        "invalid input", "data package or evidence is empty or unreadable",
        "事实包不完整或未提供", "无法生成整体判断", "未调用模型或模型结果未通过校验",
    ))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--create-template", type=Path)
    arguments = parser.parse_args()
    if arguments.create_template:
        create_template(arguments.create_template)
