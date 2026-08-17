"""Structural audit for rendered leader-facing daily report DOCX files."""

from __future__ import annotations

import zipfile
from collections.abc import Mapping
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT

from render_daily_report import ANCHORS, HIDDEN_PROJECT_NAMES


REQUIRED_HEADINGS = ("一、总体概况", "二、填报与出勤情况", "三、人员效能分析",
                     "四、项目连续性分析", "五、项目关联性及协同分析",
                     "六、风险评估", "七、管理建议")
FORBIDDEN_TOKENS = ("领导版", "待人工补充", "数据时区", "PRESALES_IN_PROGRESS",
                    "DELIVERY_IN_PROGRESS", "AFTERSALES_IN_PROGRESS", "operations-support",
                    "snapshot_captured_at", "事实包不完整或未提供",
                    "无法生成整体判断", "未调用模型或模型结果未通过校验")


@dataclass(frozen=True, slots=True)
class DocxAuditResult:
    ok: bool
    errors: tuple[str, ...]
    page_settings: Mapping[str, float]
    required_text_found: bool


def audit_daily_report(path: Path, facts: Mapping[str, object] | None = None) -> DocxAuditResult:
    errors: list[str] = []
    if not zipfile.is_zipfile(path):
        return DocxAuditResult(False, ("INVALID_DOCX_ZIP",), {}, False)
    with zipfile.ZipFile(path) as archive:
        if archive.testzip() is not None:
            errors.append("CORRUPT_DOCX_ZIP")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in _all_paragraphs(document))
    if any(anchor in text for anchor in ANCHORS):
        errors.append("UNCONSUMED_ANCHOR")
    required_text_found = all(heading in text for heading in REQUIRED_HEADINGS)
    if not required_text_found:
        errors.append("MISSING_REQUIRED_SECTION")
    if any(token in text for token in FORBIDDEN_TOKENS):
        errors.append("FORBIDDEN_PLACEHOLDER")
    if "覆盖周期：" not in text:
        errors.append("MISSING_CHINESE_PERIOD_LINE")
    section = document.sections[0]
    page_settings = {"top_cm": section.top_margin.cm, "bottom_cm": section.bottom_margin.cm,
                     "left_cm": section.left_margin.cm, "right_cm": section.right_margin.cm,
                     "width_cm": section.page_width.cm, "height_cm": section.page_height.cm}
    if (section.orientation != WD_ORIENT.PORTRAIT
            or abs(section.page_width.inches - 8.5) >= 0.02
            or abs(section.page_height.inches - 11.0) >= 0.02):
        errors.append("INVALID_PAGE_SIZE_OR_ORIENTATION")
    for margin in (section.top_margin, section.bottom_margin,
                   section.left_margin, section.right_margin):
        if abs(margin.inches - 0.83) >= 0.02:
            errors.append("INVALID_PAGE_MARGIN")
            break
    if len(document.tables) != 1 or len(document.tables[0].rows) != 2 or len(document.tables[0].columns) != 4:
        errors.append("INVALID_DASHBOARD_TABLE_STRUCTURE")
    if facts is not None:
        attendance = facts.get("attendance_summary")
        if isinstance(attendance, Mapping):
            for key in ("missing_people", "full_day_leave_people", "half_day_leave_people", "review_required_people"):
                for item in attendance.get(key, []):
                    if isinstance(item, Mapping) and isinstance(item.get("name"), str) and item["name"] not in text:
                        errors.append("MISSING_FACT_NAME")
        for section_name in ("formal_project_dynamics", "unlinked_project_dynamics", "stale_project_alerts"):
            section = facts.get(section_name)
            if isinstance(section, list):
                for item in section:
                    if isinstance(item, Mapping):
                        project_name = item.get("project_name")
                        if (isinstance(project_name, str)
                                and project_name not in HIDDEN_PROJECT_NAMES
                                and project_name not in text):
                            errors.append("MISSING_FACT_PROJECT")
        metrics = facts.get("metrics")
        status = facts.get("status_distribution")
        if isinstance(metrics, Mapping) and str(metrics.get("effective_task_count")) not in text:
            errors.append("MISSING_DASHBOARD_METRIC")
        if isinstance(status, Mapping):
            for key in ("completed", "blocked"):
                if str(status.get(key)) not in text:
                    errors.append("MISSING_DASHBOARD_METRIC")
    return DocxAuditResult(not errors, tuple(errors), page_settings, required_text_found)


def _all_paragraphs(document: Document) -> list[object]:
    return list(document.paragraphs) + [paragraph for table in document.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs]
